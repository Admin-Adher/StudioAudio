"""Exports lisibles et partageables pour les transcriptions."""

from __future__ import annotations

import csv
import json
import re
import zipfile
from dataclasses import asdict
from datetime import datetime
from pathlib import Path
from typing import Iterable, Sequence

from .core import ROOT_DIR, TranscriptTurn, TranscriptionResult


RESULTS_DIR = ROOT_DIR / "resultats"


def format_timestamp(seconds: float, *, milliseconds: bool = False) -> str:
    value = max(0.0, float(seconds))
    hours = int(value // 3600)
    minutes = int((value % 3600) // 60)
    whole_seconds = int(value % 60)
    if milliseconds:
        millis = int(round((value - int(value)) * 1000))
        if millis == 1000:
            whole_seconds += 1
            millis = 0
        return f"{hours:02d}:{minutes:02d}:{whole_seconds:02d}.{millis:03d}"
    return f"{hours:02d}:{minutes:02d}:{whole_seconds:02d}"


def parse_timestamp(value: object) -> float:
    text = str(value).strip().replace(",", ".")
    parts = text.split(":")
    try:
        if len(parts) == 3:
            return float(parts[0]) * 3600 + float(parts[1]) * 60 + float(parts[2])
        if len(parts) == 2:
            return float(parts[0]) * 60 + float(parts[1])
        return float(text)
    except (TypeError, ValueError) as exc:
        raise ValueError(f"Horodatage invalide : {value}") from exc


def format_duration(seconds: float) -> str:
    total = max(0, int(round(float(seconds))))
    hours, remainder = divmod(total, 3600)
    minutes, secs = divmod(remainder, 60)
    if hours:
        return f"{hours} h {minutes:02d} min {secs:02d} s"
    if minutes:
        return f"{minutes} min {secs:02d} s"
    return f"{secs} s"


def safe_stem(filename: str) -> str:
    stem = Path(filename).stem.strip() or "audio"
    stem = re.sub(r"[^\w.-]+", "-", stem, flags=re.UNICODE)
    stem = re.sub(r"-+", "-", stem).strip("-._")
    return (stem or "audio")[:80]


def turns_to_rows(turns: Sequence[TranscriptTurn]) -> list[list[str]]:
    return [
        [
            format_timestamp(turn.start),
            format_timestamp(turn.end),
            turn.speaker,
            turn.text,
        ]
        for turn in turns
    ]


def rows_to_turns(rows: Iterable[Sequence[object]]) -> list[TranscriptTurn]:
    turns: list[TranscriptTurn] = []
    for index, row in enumerate(rows):
        values = list(row)
        if len(values) < 4:
            continue
        text = str(values[3]).strip()
        if not text:
            continue
        start = parse_timestamp(values[0])
        end = max(start, parse_timestamp(values[1]))
        speaker = str(values[2]).strip() or "Interlocuteur"
        turns.append(
            TranscriptTurn(
                start=start,
                end=end,
                speaker_id=f"EDITED_{index:03d}",
                speaker=speaker,
                text=text,
            )
        )
    return turns


def plain_transcript(turns: Sequence[TranscriptTurn]) -> str:
    blocks = [
        f"[{format_timestamp(turn.start)}] {turn.speaker}\n{turn.text}"
        for turn in turns
    ]
    return "\n\n".join(blocks).strip() + ("\n" if blocks else "")


def _srt_timestamp(seconds: float) -> str:
    return format_timestamp(seconds, milliseconds=True).replace(".", ",")


def _vtt_timestamp(seconds: float) -> str:
    return format_timestamp(seconds, milliseconds=True)


def _write_docx(path: Path, result: TranscriptionResult) -> None:
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.shared import Pt

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    document.add_heading("Transcription", level=0)
    metadata = document.add_paragraph()
    metadata.add_run("Fichier : ").bold = True
    metadata.add_run(result.audio_name)
    metadata.add_run("\nDurée : ").bold = True
    metadata.add_run(format_duration(result.duration))
    metadata.add_run("\nLangue : ").bold = True
    metadata.add_run(result.language)
    metadata.add_run("\nProfil : ").bold = True
    metadata.add_run(result.profile)

    document.add_paragraph()
    for turn in result.turns:
        paragraph = document.add_paragraph()
        header = paragraph.add_run(
            f"[{format_timestamp(turn.start)}] {turn.speaker}"
        )
        header.bold = True
        header.add_break(WD_BREAK.LINE)
        paragraph.add_run(turn.text)

    document.save(str(path))


def _write_files(directory: Path, result: TranscriptionResult) -> list[Path]:
    directory.mkdir(parents=True, exist_ok=True)
    base = safe_stem(result.audio_name) + "-transcription"

    txt_path = directory / f"{base}.txt"
    docx_path = directory / f"{base}.docx"
    csv_path = directory / f"{base}.csv"
    srt_path = directory / f"{base}.srt"
    vtt_path = directory / f"{base}.vtt"
    json_path = directory / f"{base}.json"

    txt_path.write_text(plain_transcript(result.turns), encoding="utf-8-sig")

    with csv_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle, delimiter=";")
        writer.writerow(["Début", "Fin", "Interlocuteur", "Texte"])
        writer.writerows(turns_to_rows(result.turns))

    srt_blocks: list[str] = []
    for index, turn in enumerate(result.turns, start=1):
        srt_blocks.append(
            f"{index}\n{_srt_timestamp(turn.start)} --> {_srt_timestamp(turn.end)}\n"
            f"{turn.speaker} : {turn.text}"
        )
    srt_path.write_text("\n\n".join(srt_blocks) + "\n", encoding="utf-8-sig")

    vtt_lines = ["WEBVTT", ""]
    for turn in result.turns:
        vtt_lines.extend(
            [
                f"{_vtt_timestamp(turn.start)} --> {_vtt_timestamp(turn.end)}",
                f"{turn.speaker} : {turn.text}",
                "",
            ]
        )
    vtt_path.write_text("\n".join(vtt_lines), encoding="utf-8-sig")

    json_path.write_text(
        json.dumps(result.to_dict(), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    _write_docx(docx_path, result)

    return [docx_path, txt_path, csv_path, srt_path, vtt_path, json_path]


def write_exports(result: TranscriptionResult, *, corrected: bool = False) -> list[Path]:
    RESULTS_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S-%f")[:-3]
    suffix = "-corrigee" if corrected else ""
    folder = RESULTS_DIR / f"{safe_stem(result.audio_name)}-{timestamp}{suffix}"
    files = _write_files(folder, result)

    zip_path = folder / f"{safe_stem(result.audio_name)}-transcription{suffix}.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in files:
            archive.write(path, arcname=path.name)
    return [zip_path, *files]


def write_batch_archive(zip_paths: Sequence[Path]) -> Path:
    """Regroupe les ZIP individuels d'une file d'attente dans un seul ZIP."""
    archives = [Path(path) for path in zip_paths if Path(path).is_file()]
    if not archives:
        raise ValueError("Aucune transcription n'est disponible pour créer le lot.")

    RESULTS_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S-%f")[:-3]
    folder = RESULTS_DIR / f"lot-{timestamp}"
    folder.mkdir(parents=True, exist_ok=True)
    batch_path = folder / "lot-transcriptions.zip"

    with zipfile.ZipFile(batch_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for index, path in enumerate(archives, start=1):
            archive.write(path, arcname=f"{index:02d}-{path.name}")
    return batch_path


def result_from_rows(
    rows: Iterable[Sequence[object]],
    metadata: dict[str, object],
) -> TranscriptionResult:
    turns = rows_to_turns(rows)
    if not turns:
        raise ValueError("Le tableau corrigé ne contient aucun texte à exporter.")
    return TranscriptionResult(
        audio_name=str(metadata.get("audio_name") or "audio"),
        duration=float(metadata.get("duration") or turns[-1].end),
        processing_seconds=float(metadata.get("processing_seconds") or 0.0),
        profile=str(metadata.get("profile") or "Équilibré"),
        model=str(metadata.get("model") or ""),
        language=str(metadata.get("language") or "inconnue"),
        language_probability=(
            float(metadata["language_probability"])
            if metadata.get("language_probability") is not None
            else None
        ),
        requested_speakers=int(metadata.get("requested_speakers") or 1),
        detected_speakers=len({turn.speaker for turn in turns}),
        turns=turns,
        warnings=[str(item) for item in metadata.get("warnings", [])],
    )
